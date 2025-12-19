import os
import argparse
from pathlib import Path
from typing import List
import chromadb
import docx 
from llama_index.core import Settings, StorageContext, VectorStoreIndex, Document
from llama_index.vector_stores.chroma import ChromaVectorStore
from llama_index.core.node_parser import HierarchicalNodeParser, get_leaf_nodes
from core import Embedding

Settings.context_window = 32000
Settings.num_output = 4000

def read_docx_text(file_path: Path) -> str:
    try:
        doc = docx.Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        return "\n\n".join(full_text)
    except Exception as e:
        print(f"Lỗi đọc nội dung file Word {file_path.name}: {e}")
        return ""

def load_word_files(folder_path: str, dataset_tag: str) -> List[Document]:
    p = Path(folder_path)
    if not p.exists():
        raise FileNotFoundError(f"Không tìm thấy thư mục: {folder_path}")
    
    word_files = list(p.rglob("*.docx"))
    docs = []
    for file_path in word_files:
        if file_path.name.startswith("~$"):
            continue
        content = read_docx_text(file_path)
        if not content:
            continue
        title = file_path.stem 
        meta = {
            "source": file_path.name,
            "domain": "word_import",
            "url": file_path.name,
            "title": title,
            "dataset": dataset_tag, 
        }
        full_text = f"# {title}\n\n{content}"
        docs.append(Document(text=full_text, metadata=meta))
    return docs

def main():
    parser = argparse.ArgumentParser(description="")
    
    parser.add_argument("--persist-dir", type=str, default="./chroma_store")
    parser.add_argument("--collection", type=str, default="emb")
    parser.add_argument("--docstore-dir", type=str, default="./docstore_save")
    parser.add_argument("--folder", type=str, required=True, help="")
    parser.add_argument("--tag", type=str, default="tai_lieu_word", help="")
    parser.add_argument("--chunk-sizes", type=int, nargs="+", default=[1024, 512])
    parser.add_argument("--chunk-overlap", type=int, default=128)

    args = parser.parse_args()

    Settings.embed_model = Embedding(embed_batch_size=5) 
    documents = load_word_files(args.folder, dataset_tag=args.tag)
    
    if not documents:
        raise RuntimeError("Không có document nào để xử lý.")

    node_parser = HierarchicalNodeParser.from_defaults(
        chunk_sizes=args.chunk_sizes,
        chunk_overlap=args.chunk_overlap
    )

    all_nodes = node_parser.get_nodes_from_documents(documents)
    leaf_nodes = get_leaf_nodes(all_nodes)
    db = chromadb.PersistentClient(path=args.persist_dir)
    chroma_collection = db.get_or_create_collection(args.collection)
    vector_store = ChromaVectorStore(chroma_collection=chroma_collection)

    docstore_path = os.path.join(args.docstore_dir, "docstore.json")
    
    if os.path.exists(docstore_path):
        storage_context = StorageContext.from_defaults(
            persist_dir=args.docstore_dir, 
            vector_store=vector_store
        )
    else:
        storage_context = StorageContext.from_defaults(vector_store=vector_store)
    storage_context.docstore.add_documents(all_nodes)
    VectorStoreIndex(
        leaf_nodes,
        storage_context=storage_context,
        show_progress=True
    )
    storage_context.persist(persist_dir=args.docstore_dir)
    print("HOÀN TẤT!")

if __name__ == "__main__":
    main()